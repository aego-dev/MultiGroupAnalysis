import io
import re
import pandas as pd
import numpy as np
from scipy import stats
from scipy.stats import chi2_contingency, fisher_exact, kruskal
from itertools import combinations
from lifelines import CoxPHFitter, KaplanMeierFitter
from lifelines.statistics import logrank_test as lr_test, multivariate_logrank_test
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import streamlit as st

matplotlib.rcParams.update({"font.family": "DejaVu Sans", "font.size": 11})

st.set_page_config(page_title="Multi-Group Analysis", layout="wide")
st.title("Multi-Group Analysis")

# ═════════════════════════════════════════════════════════════════════════════
# STEP 0 — Number of groups  (shown before file upload)
# ═════════════════════════════════════════════════════════════════════════════
st.header("0. How Many Groups Are You Analyzing?")
num_groups = st.radio(
    "Select the number of groups:",
    options=[2, 3, 4],
    horizontal=True,
)

# ═════════════════════════════════════════════════════════════════════════════
# STEP 1 — File upload
# ═════════════════════════════════════════════════════════════════════════════
st.header("1. Upload Your Data")
uploaded = st.file_uploader("Upload your Excel file", type=["xlsx", "xls"])
if not uploaded:
    st.info("Upload an Excel file to continue.")
    st.stop()

df = pd.read_excel(uploaded)
st.success(f"Loaded {len(df):,} rows and {len(df.columns)} columns.")
col_options = df.columns.tolist()

# ── Reset Cox covariate session_state when the uploaded file changes ──────────
_file_fp = f"{uploaded.name}_{len(df)}_{len(col_options)}"
if st.session_state.get("_file_fp") != _file_fp:
    keys_to_clear = [k for k in st.session_state
                     if k.startswith("cox_sel_") or k.startswith("cox_search_")
                     or k.startswith("cox_sdur_") or k.startswith("cox_sevt_")]
    for k in keys_to_clear:
        del st.session_state[k]
    st.session_state["_file_fp"] = _file_fp

# ── Fuzzy column matcher ──────────────────────────────────────────────────────
def normalize(s):
    return re.sub(r"[^a-z0-9]", "", str(s).lower())

def find_column(df, keywords):
    norm_cols = {col: normalize(col) for col in df.columns}
    for kw in keywords:
        nkw = normalize(kw)
        for col, ncol in norm_cols.items():
            if nkw in ncol or ncol in nkw:
                return col
    return None

def resolve_vars(df, var_dict):
    resolved = {}
    for label, keywords in var_dict.items():
        col = find_column(df, keywords)
        if col:
            resolved[label] = col
    return resolved

def searchable_selectbox(label, options, key, default=None):
    """A selectbox with a live search filter above it."""
    search = st.text_input(f"🔍 Search for {label}", value="", key=f"search_sel_{key}",
                           placeholder="Type to filter columns...")
    filtered = [o for o in options if search.lower() in o.lower()] if search else options
    if not filtered:
        st.warning("No columns match — showing all.")
        filtered = options
    default_idx = filtered.index(default) if default in filtered else 0
    return st.selectbox(label, filtered, index=default_idx, key=key)

# ── Default variable definitions ──────────────────────────────────────────────
DEFAULT_CONT_BASE = {
    "Age":                          ["Age", "Pat_Age_at_Encounter"],
    "BMI":                          ["BMI"],
    "STS Risk Score":               ["STSRiskScore", "STS_risk_score"],
    "Preprocedural Hematocrit":     ["Preop_Lab_Hematocrit", "Preop_Hematocrit"],
    "Preprocedural Hemoglobin":     ["Preop_Hemoglobin", "Hb"],
    "Preprocedural Creatinine":     ["Preop_Creatinine", "Creatinine"],
    "eGFR":                         ["eGFR"],
    "Systolic Blood Pressure, mm Hg":  ["Systolic_bp_pre", "sbp_pre", "systolic_bp"],
    "Diastolic Blood Pressure, mm Hg": ["Diastolic_bp_pre", "dbp_pre", "diastolic_bp"],
}
DEFAULT_CAT_BASE = {
    # Demographics
    "Female":                           ["Female", "Gender"],
    "White":                            ["Race_white=1,black=2,other=3", "white_race", "race_white"],
    "Black":                            ["Race_white=1,black=2,other=3", "black_race", "race_black"],
    "Other Race":                       ["Race_white=1,black=2,other=3", "other_race"],
    # History
    "Previous MV Surgery":              ["PrevProcMV_Surgery", "HX_MV_Surgery"],
    "Coronary Artery Disease":          ["CAD", "HX_CAD"],
    "Prior PCI":                        ["PriorPCI", "HX_PCI"],
    "Prior CABG":                       ["PriorCABG", "HX_CABG"],
    "Known Left Main Disease":          ["Known_LMainDis", "HX_Left_Main"],
    "Prior Myocardial Infarction":      ["PriorMI", "HX_Myocardial_Iarction", "Prior MI"],
    "Prior Stroke":                     ["PriorStroke", "HX_Stroke", "Prior Stroke"],
    "Transient Ischemic Attack":        ["TIA", "HX_TIA"],
    "Carotid Artery Disease":           ["Carotid_disease", "HX_Carotid_Artery_Disease"],
    "Peripheral Artery Disease":        ["PriorPAD", "HX_PAD"],
    "Hypertension":                     ["Hypertension", "HX_Hypertension", "HTN"],
    "Diabetes Mellitus":                ["Diabetes", "HX_Diabetes", "DM"],
    "Hyperlipidemia":                   ["Hyperlipidemia", "HDL", "HDM", "hyperlipidemia"],
    "Cancer":                           ["Cancer", "HX_Cancer", "cancer"],
    "GI Bleeding":                      ["GI Bleeding", "GI_Bleeding", "gi_bleeding"],
    "eGFR < 25":                        ["eGFR <25", "eGFR_less_25", "egfr_lt25"],
    "Current/Recent Smoker":            ["CurrentRecent_Smoker", "smoking", "smoker"],
    "Dialysis":                         ["CurrentDialysis", "Preop_Renal_Dialysis"],
    "Chronic Lung Disease":             ["ChrLungD", "HX_Chronic_Lung_Disease"],
    "NYHA Class III or IV":             ["NYHAIIIorIV", "Preop_NYHA_3_4", "NYHA class"],
    "Prior Cardiac Shock":              ["PriorCardioShock", "HX_Cardiogenic_Shock"],
    "Porcelain Aorta":                  ["PorcelainAorta", "porcelain_aorta"],
    "Atrial Fibrillation or Flutter":   ["History_Afib_Afluter", "HX_Afib", "AF/Af"],
    "Prior ICD":                        ["PrevICD", "HX_ICD_Implant"],
    "Pacemaker/ICD":                    ["Pacer/ICD", "Pacemaker_ICD", "pacer_icd", "pacemaker_icd"],
    # Procedural
    "Predilation":                      ["Predilation", "predilation"],
    "Postdilation":                     ["Postdilation", "postdilation"],
    "Balloon-Expanding Valve":          ["Ballon_1_Self_2", "Valve_Type_Balloon1_Self2_Lotus3"],
    "Self-Expanding Valve":             ["Ballon_1_Self_2", "Valve_Type_Balloon1_Self2_Lotus3"],
}
DEFAULT_ECHO = {
    "Heart Rate, beats/min":                        ["Heart Rate", "heart_rate", "hr_bpm"],
    "Systolic Blood Pressure, mm Hg":               ["Blood Pressure", "sbp", "systolic_bp"],
    "LVEF, %":                                      ["LVEF", "LVEF_pre", "ejectionfraction"],
    "LV End-Diastolic Volume, mL":                  ["LV_end_diastolic_volume", "LVEDV", "lvedv"],
    "LV End-Systolic Volume, mL":                   ["LV_end_systolic_volume", "LVESV", "lvesv"],
    "Maximal Aortic Size, cm":                      ["AO_Size_Max", "max_aortic_size", "aortic_size"],
    "AV Area, cm²":                                 ["AV_Area", "AV_area_pre", "avarea"],
    "LA Volume, mL":                                ["LA_Volume", "LA_volume_pre", "lavolume"],
    "LA Diameter, cm":                              ["LA_Diameter", "LA_diameter_pre", "ladiameter"],
    "AV Mean Gradient, mm Hg":                      ["AV_Mean_Gradient", "AV_LVOT_Mean Gradient", "avmeangradient"],
    "AV Peak Gradient, mm Hg":                      ["AV_Peak Gradient ", "AV_LVOT_Peak_Gradient", "avpeakgradient"],
    "LVOT Diameter, cm":                            ["LVOT Diameter", "LVOT_diameter", "lvotdiameter"],
    "LVOT Maximum Velocity, m/s":                   ["LVOT_Velocity_Max", "LVOT_max_velocity"],
    "LVOT Stroke Volume, mL":                       ["LVOT_Stroke_Volume", "LVOT_stroke_volume"],
    "Aortic Valve Resistance":                      ["Aortic Valve Resistance (h * HR * SEP/CO *1.33), h is mean aortic valve gradient", "avresistance"],
    # Echocardiographic hemodynamics
    "Inferior Vena Cava, cm":                       ["IVC", "ivc", "inferior_vena_cava"],
    "Right Atrial Pressure, mm Hg":                 ["RAP", "rap", "right_atrial_pressure"],
    "Right Ventricular Systolic Pressure, mm Hg":   ["RVSP", "rvsp", "right_ventricular_systolic_pressure"],
    "Mitral Regurgitation Grade":                   ["MR", "mr", "mitral_regurgitation"],
    "Aortic Regurgitation Grade":                   ["AR", "ar", "aortic_regurgitation"],
    # Right heart catheterization
    "Right Atrial V-Wave, mm Hg":                   ["RA_v_wave_pre", "ra_v_wave"],
    "Right Atrial Mean Pressure, mm Hg":            ["RA_mean_pre", "ra_mean"],
    "Pulmonary Capillary Wedge V-Wave, mm Hg":      ["PCW_v_wave_pre", "pcw_v_wave"],
    "Pulmonary Capillary Wedge Mean Pressure, mm Hg": ["PCWP_mean_pre", "pcwp_mean"],
    "Right Ventricular Systolic Pressure (RHC), mm Hg": ["RV_pressure_S", "rv_pressure_s"],
    "Right Ventricular Diastolic Pressure, mm Hg":  ["RV_pressure_D", "rv_pressure_d"],
    "Right Ventricular End-Diastolic Pressure, mm Hg": ["RV_pressure_E", "rv_pressure_e"],
    "Cardiac Output, L/min":                        ["Cardiac_output_pre_fick", "cardiac_output"],
    "Cardiac Index, L/min/m²":                      ["Cardiac_index_pre_fick", "cardiac_index"],
    # Doppler
    "Pre AT":                                       ["Pre_AV_AT(Delta T)", "Pre_AT", "preat"],
    "Pre ET":                                       ["Pre_AV_ET (Delta T)", "Pre_ET", "preet"],
    "Pre AT/ET":                                    ["Pre_AV_AT_ET_Ratio", "Pre_AT_ET", "preatеtratio"],
    "Post AT":                                      ["Post_AV_AT(Delta T)", "Post_AT", "postat"],
    "Post ET":                                      ["Post_AV_ET (Delta T)", "Post_ET", "postet"],
    "Post AT/ET":                                   ["Post_AV_AT_ET_Ratio", "Post_AT_ET", "postatetрatio"],
}
DEFAULT_OUTCOMES_CAT = {
    # 30-day binary outcomes
    "LBBB 30 Days":                      ["GROUP_No_LBBB_0_LBBB_After_1_LBBB_Before_2", "LBBB 30 days", "lbbb30"],
    "PPM < 30 Days":                     ["PPM_in_Less_than_30_Days", "PPM 30 days", "ppm30"],
    "PPM Implantation":                  ["PPM implantation", "PPM_implantation", "ppm_implantation"],
    "ICD Implantation":                  ["ICD implantation", "ICD_implantation", "icd_implantation"],
    "Atrial Fibrillation Admission":     ["Afib admission", "Afib_admission", "afib_admission"],
    "Bleeding":                          ["Bleeding", "bleeding"],
    # Pulmonary hypertension categories
    "Pulmonary Hypertension":            ["PH (Moderate-to-severe: RVSP >45)", "PH", "pulmonary_hypertension"],
    "Moderate-to-Severe Pulmonary Hypertension (RVSP > 45 mm Hg)": ["PH (Moderate-to-severe: RVSP >45)", "ph_moderate_severe"],
    "Severe Pulmonary Hypertension (RVSP > 55 mm Hg)":             ["Severe PH RVSP>55", "ph_severe"],
    # Table 3 categorical
    "Sex (Female)":                      ["Female", "Gender"],
    "Diabetes Mellitus":                 ["Diabetes", "HX_Diabetes", "DM"],
    "Chronic Lung Disease":              ["ChrLungD", "HX_Chronic_Lung_Disease"],
    "Hypertension":                      ["Hypertension", "HX_Hypertension", "HTN"],
    "Coronary Artery Disease":           ["CAD", "HX_CAD"],
    "Predilation":                       ["Predilation", "predilation"],
}
DEFAULT_OUTCOMES_CONT = {
    # Post-TAVR echo
    "Post-TAVR LVEF, %":                 ["LVEF", "LVEF_post", "lvef_post"],
    "Post-TAVR LV End-Diastolic Volume, mL": ["LV_end_diastolic_volume", "lvedv_post"],
    "Post-TAVR LV End-Systolic Volume, mL":  ["LV_end_systolic_volume", "lvesv_post"],
    "Post-TAVR AV Mean Gradient, mm Hg":     ["AV_Mean_Gradient", "avmeangradient_post"],
    "Post-TAVR AV Peak Gradient, mm Hg":     ["AV_Peak Gradient ", "avpeakgradient_post"],
    # AT/ET
    "Pre-TAVR AT/ET Ratio":              ["Pre_AV_AT_ET_Ratio", "Pre_AT_ET"],
    "Post-TAVR AT/ET Ratio":             ["Post_AV_AT_ET_Ratio", "Post_AT_ET"],
    # Other continuous outcomes
    "Age":                               ["Age", "Pat_Age_at_Encounter"],
    "BMI":                               ["BMI"],
    "LVEF, %":                           ["LVEF", "LVEF_pre"],
    "LV End-Diastolic Volume, mL":       ["LV_end_diastolic_volume", "lvedv"],
    "LV End-Systolic Volume, mL":        ["LV_end_systolic_volume", "lvesv"],
    "AV Mean Gradient, mm Hg":           ["AV_Mean_Gradient", "AV_LVOT_Mean Gradient"],
    "AV Peak Gradient, mm Hg":           ["AV_Peak Gradient ", "AV_LVOT_Peak_Gradient"],
    "Pre-TAVR AT":                       ["Pre_AV_AT(Delta T)", "Pre_AT"],
    "Post-TAVR AT":                      ["Post_AV_AT(Delta T)", "Post_AT"],
}
# Keep a combined alias for any legacy references
DEFAULT_OUTCOMES = {**DEFAULT_OUTCOMES_CAT, **DEFAULT_OUTCOMES_CONT}

GROUP_COLORS = ["#d62728", "#1f77b4", "#2ca02c", "#ff7f0e"]
GROUP_FC     = ["#FDECEA", "#E8F4FD", "#EAF7EA", "#FFF3E0"]
GROUP_EC     = ["#d62728", "#1f77b4", "#27ae60", "#ff7f0e"]

# Section-header rows to inject into Table 1 (baseline) — value is the label shown before the group
T1_SECTION_HEADERS = {
    "White":                            "Race",
    "Previous MV Surgery":              "History",
    "eGFR":                             "Renal Function",
    "Pacemaker/ICD":                    "Devices",
    "Predilation":                      "Procedural",
    "Preprocedural Hematocrit":         "Preprocedural Labs",
    "Systolic Blood Pressure, mm Hg":   "Preprocedural Vitals",
}

# Section-header rows for Table 2 (echo)
T2_SECTION_HEADERS = {
    "LV End-Diastolic Volume, mL":                      "Left Ventricular Function",
    "Inferior Vena Cava, cm":                           "Right Heart Parameters",
    "Right Atrial V-Wave, mm Hg":                       "Right Heart Catheterization",
    "Pre AT":                                           "Pre-Procedure Doppler",
    "Post AT":                                          "Post-Procedure Doppler",
}

# Section-header rows for Table 3 (outcomes)
T3_SECTION_HEADERS = {
    "PPM Implantation":             "Clinical Events",
    "Pulmonary Hypertension":       "Pulmonary Hypertension",
    "Post-TAVR LVEF, %":            "Post-Procedure Echo",
    "Pre-TAVR AT/ET Ratio":         "AT/ET Ratios",
}

# ═════════════════════════════════════════════════════════════════════════════
# STEP 2 — Study title
# ═════════════════════════════════════════════════════════════════════════════
st.header("2. Study Title")
default_title = f"{num_groups}-Group Analysis"
study_title = st.text_input(
    "Enter your study title (used as the document title and filename):",
    value=default_title,
)
safe_filename = re.sub(r"[^\w\s-]", "", study_title).strip().replace(" ", "_") or f"{num_groups}_Group_Analysis"

# ═════════════════════════════════════════════════════════════════════════════
# STEP 3 — Define groups
# ═════════════════════════════════════════════════════════════════════════════
st.header("3. Define Your Groups")

group_col = st.selectbox(
    f"Select the group column (must contain values 0\u2013{num_groups - 1}):",
    col_options,
    index=col_options.index("Group") if "Group" in col_options else 0,
)

valid_vals = set(range(num_groups))
found_vals = set(df[group_col].dropna().unique())
if found_vals - valid_vals:
    st.warning(f"Group column contains unexpected values. Expected only: {sorted(valid_vals)}")

label_cols = st.columns(num_groups)
labels = []
for v in range(num_groups):
    with label_cols[v]:
        st.markdown(f"**Group {v} (value = {v})**")
        lbl = st.text_input("Label", value=f"Group {v}", key=f"g{v}")
        labels.append(lbl)

# ═════════════════════════════════════════════════════════════════════════════
# STAT HELPERS
# ═════════════════════════════════════════════════════════════════════════════
def fmt_p(p):
    if pd.isna(p):
        return "N/A"
    return "<0.001" if p < 0.001 else f"{p:.3f}"

def is_significant(p_str):
    if p_str == "<0.001":
        return True
    if p_str == "N/A":
        return False
    try:
        return float(p_str) < 0.05
    except (ValueError, TypeError):
        return False

def calc_continuous(groups):
    """Returns list of (mean, std) per group and an omnibus p-value."""
    # Coerce to numeric — silently drops strings/non-numeric values
    cleaned = [pd.to_numeric(g, errors="coerce").dropna() for g in groups]
    stats_per_group = [(float(np.mean(g)), float(np.std(g, ddof=1))) if len(g) > 0 else (float("nan"), float("nan")) for g in cleaned]
    try:
        if len(cleaned) == 2:
            _, p = stats.ttest_ind(cleaned[0], cleaned[1], nan_policy="omit")
        elif any(len(g) < 3 for g in cleaned):
            _, p = kruskal(*cleaned)
        else:
            _, p = stats.f_oneway(*cleaned)
    except Exception:
        p = float("nan")
    return stats_per_group, p

# Race value map: label -> (column_keywords, race_code_value)
# When a categorical label maps to a coded column (1=white,2=black,3=other),
# the dict below tells build_stats_df which numeric value means "present".
RACE_VALUE_MAP = {
    "White":      1,
    "Black":      2,
    "Other Race": 3,
    # Valve type (1=balloon, 2=self-expanding)
    "Balloon-Expanding Valve": 1,
    "Self-Expanding Valve":    2,
}

def calc_categorical(groups, race_value=None):
    """Returns list of (count, pct) per group and an omnibus p-value.
    If race_value is given, counts rows equal to that value instead of summing."""
    cleaned = [pd.to_numeric(g, errors="coerce").dropna() for g in groups]
    cp = []
    for g in cleaned:
        total = len(g)
        if race_value is not None:
            count = int((g == race_value).sum())
        else:
            # Standard binary column: values should be 0/1
            # Guard against coded columns by clipping to [0,1]
            count = int(np.sum(g.clip(0, 1).round()))
        pct = count / total * 100 if total > 0 else 0.0
        cp.append((count, pct, total))
    # Build contingency table
    table = np.array([[c[0], c[2] - c[0]] for c in cp])
    try:
        if len(groups) == 2:
            if np.min(table) < 5:
                _, p = fisher_exact(table)
            else:
                _, p, _, _ = chi2_contingency(table)
        else:
            _, p, _, _ = chi2_contingency(table)
    except Exception:
        p = float("nan")
    return [(c[0], c[1]) for c in cp], p

def build_stats_df(df, group_col, variables, var_type, labels):
    n = num_groups
    ns = [(df[group_col] == v).sum() for v in range(n)]
    headers = [f"{labels[v]} (N={ns[v]})" for v in range(n)]
    rows = []
    for label, col in variables.items():
        grp_data = [df[df[group_col] == v][col].dropna() for v in range(n)]
        if any(len(g) == 0 for g in grp_data):
            continue
        if var_type == "continuous":
            sp, p = calc_continuous(grp_data)
            row = {"Characteristic": label}
            for i, (mean, std) in enumerate(sp):
                row[headers[i]] = f"{mean:.2f} \u00b1 {std:.2f}"
            row["P-Value"] = fmt_p(p)
        else:
            race_value = RACE_VALUE_MAP.get(label, None)
            cp, p = calc_categorical(grp_data, race_value=race_value)
            row = {"Characteristic": label}
            for i, (count, pct) in enumerate(cp):
                row[headers[i]] = f"{count} ({pct:.1f}%)"
            row["P-Value"] = fmt_p(p)
        rows.append(row)
    return pd.DataFrame(rows)

def inject_section_headers(df_table, header_map):
    """Insert bold section-header rows before specified variable rows."""
    if not header_map:
        return df_table
    rows_out = []
    for _, row in df_table.iterrows():
        char = row.get("Characteristic", "")
        if char in header_map:
            blank = {c: "" for c in df_table.columns}
            blank["Characteristic"] = f"  {header_map[char]}"
            blank["_section_header"] = True
            rows_out.append(blank)
        r = row.to_dict()
        r.setdefault("_section_header", False)
        rows_out.append(r)
    return pd.DataFrame(rows_out)

def style_pvalues(df):
    display_df = df[[c for c in df.columns if c != "_section_header"]].copy()
    is_hdr = df.get("_section_header", pd.Series(False, index=df.index))

    def style_row(row):
        idx = row.name
        if is_hdr.iloc[idx] if hasattr(is_hdr, 'iloc') else False:
            return ["font-weight: bold"] * len(row)
        styles = []
        for col, val in row.items():
            if col == "P-Value" and is_significant(str(val)):
                styles.append("font-weight: bold; color: #d62728")
            else:
                styles.append("")
        return styles

    return display_df.style.apply(style_row, axis=1)

def write_docx_table(doc, heading, df_table):
    doc.add_heading(heading, level=2)
    display_cols = [c for c in df_table.columns if c != "_section_header"]
    n_cols = len(display_cols)
    table = doc.add_table(rows=1, cols=n_cols)
    table.style = "Light Grid Accent 1"

    # Header row
    for i, col in enumerate(display_cols):
        run = table.rows[0].cells[i].paragraphs[0].add_run(col)
        run.bold = True

    for _, row in df_table.iterrows():
        is_header = bool(row.get("_section_header", False))
        cells = table.add_row().cells

        if is_header:
            # Merge all cells in this row into one spanning cell
            merged = cells[0]
            for i in range(1, n_cols):
                merged = merged.merge(cells[i])
            # Write the section label text, bold only (no background color)
            merged.paragraphs[0].clear()
            run = merged.paragraphs[0].add_run(str(row.get("Characteristic", "")).strip())
            run.bold = True
        else:
            for i, col_name in enumerate(display_cols):
                val = str(row.get(col_name, ""))
                run = cells[i].paragraphs[0].add_run(val)
                if col_name == "P-Value" and is_significant(val):
                    run.bold = True

    doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
# PLOT HELPERS
# ═════════════════════════════════════════════════════════════════════════════
def plot_study_flow(df, group_col, labels, study_title):
    n = num_groups
    total    = len(df)
    ns       = [(df[group_col] == v).sum() for v in range(n)]
    excluded = total - sum(ns)

    # ── Responsive sizing based on group count ────────────────────────────────
    # Each group needs enough horizontal space; min box width drives everything
    MIN_BOX_W  = 2.8          # minimum box width in data units
    H_MARGIN   = 1.2          # left/right margin
    BOX_GAP    = 0.35         # minimum gap between adjacent group boxes

    # Figure width: enough to fit n boxes + gaps + margins
    fig_w  = max(10, n * MIN_BOX_W + (n - 1) * BOX_GAP + 2 * H_MARGIN)
    fig_h  = 8.5              # fixed height — vertical layout is the same for all n

    fig, ax = plt.subplots(figsize=(fig_w, fig_h))
    xlim = fig_w              # data units == inches so layout math stays simple
    ylim = 11.0
    ax.set_xlim(0, xlim)
    ax.set_ylim(0, ylim)
    ax.axis("off")

    cx = xlim / 2.0

    # Evenly space group boxes across the available width
    usable_w  = xlim - 2 * H_MARGIN
    spacing   = usable_w / (n - 1) if n > 1 else 0
    x_pos     = [H_MARGIN + i * spacing for i in range(n)]

    # Box width: fill ~85 % of the spacing slot, but never exceed MIN_BOX_W * 1.6
    box_w     = min(spacing * 0.82, MIN_BOX_W * 1.6) if n > 1 else MIN_BOX_W * 1.4

    # Font sizes scale down slightly for 4 groups to keep text from clipping
    lbl_fs    = 9 if n <= 2 else (8.5 if n == 3 else 7.5)
    hdr_fs    = 10 if n <= 2 else 9

    # Box height — slightly taller for more breathing room
    BOX_H     = 1.0

    # ── Helper functions ──────────────────────────────────────────────────────
    def box(x, y, w, h, text, fc="#E8F4FD", ec="#2c7bb6", fontsize=9):
        rect = mpatches.FancyBboxPatch(
            (x - w / 2, y - h / 2), w, h,
            boxstyle="round,pad=0.15",
            facecolor=fc, edgecolor=ec, linewidth=1.5,
            clip_on=False,
        )
        ax.add_patch(rect)
        ax.text(x, y, text, ha="center", va="center", fontsize=fontsize,
                multialignment="center", clip_on=False)

    def arrow(x, y_start, y_end):
        ax.annotate("", xy=(x, y_end), xytext=(x, y_start),
                    arrowprops=dict(arrowstyle="->", color="black", lw=1.5))

    def n_way_fork(x_center, y_start, x_positions, y_end):
        """Draw a stem-and-branches fork from one point to many."""
        y_mid = (y_start + y_end) / 2.0
        ax.plot([x_center, x_center], [y_start, y_mid], color="black", lw=1.5)
        ax.plot([min(x_positions), max(x_positions)], [y_mid, y_mid],
                color="black", lw=1.5)
        for xp in x_positions:
            ax.annotate("", xy=(xp, y_end), xytext=(xp, y_mid),
                        arrowprops=dict(arrowstyle="->", color="black", lw=1.5))

    # ── Vertical positions ────────────────────────────────────────────────────
    # Groups are the final endpoint — no analysis box below them
    if excluded > 0:
        y_total    = 10.0
        y_excluded =  8.2
        y_included =  6.4
        y_groups   =  4.2
    else:
        y_total    = 10.0
        y_groups   =  6.8

    half_h = BOX_H / 2.0

    # ── Draw top boxes ────────────────────────────────────────────────────────
    box(cx, y_total, 5.5, BOX_H,
        f"Total patients enrolled\nN = {total:,}",
        fontsize=hdr_fs)

    if excluded > 0:
        arrow(cx, y_total - half_h, y_excluded + half_h)
        excl_vals = ", ".join(str(v) for v in range(n))
        box(cx, y_excluded, 6.0, BOX_H,
            f"Excluded (not in groups {excl_vals})\nn = {excluded:,}",
            fc="#FFF3CD", ec="#FFC107", fontsize=lbl_fs)
        arrow(cx, y_excluded - half_h, y_included + half_h)
        box(cx, y_included, 5.5, BOX_H,
            f"Included in analysis\nN = {sum(ns):,}",
            fontsize=hdr_fs)
        n_way_fork(cx, y_included - half_h, x_pos, y_groups + half_h)
    else:
        n_way_fork(cx, y_total - half_h, x_pos, y_groups + half_h)

    # ── Group boxes — these are the final endpoint, no arrows below ───────────
    for i in range(n):
        box(x_pos[i], y_groups, box_w, BOX_H,
            f"{labels[i]}\nn = {ns[i]:,}",
            fc=GROUP_FC[i], ec=GROUP_EC[i], fontsize=lbl_fs)

    plt.tight_layout()
    return fig


def plot_km(df, group_col, time_col, event_col, labels, fig_label, risk_interval=200):
    n = num_groups
    THREE_YEARS = 1095  # days

    arms = {}
    for v in range(n):
        sub = df[df[group_col] == v].copy()
        t = pd.to_numeric(sub[time_col], errors="coerce")
        e = pd.to_numeric(sub[event_col], errors="coerce")
        idx = t.dropna().index.intersection(e.dropna().index)
        t = t.loc[idx]
        e = e.loc[idx]
        # Cap at 3 years; censor anything beyond
        e = e.where(t <= THREE_YEARS, other=0)
        t = t.clip(upper=THREE_YEARS)
        arms[v] = (t, e)

    if n == 2:
        lr = lr_test(arms[0][0], arms[1][0],
                     event_observed_A=arms[0][1], event_observed_B=arms[1][1])
        p_global     = lr.p_value
        p_global_str = "<0.001" if p_global < 0.001 else f"{p_global:.3f}"
        p_text       = f"Log-rank p = {p_global_str}"
    else:
        all_t = pd.concat([arms[v][0] for v in range(n)])
        all_e = pd.concat([arms[v][1] for v in range(n)])
        all_g = pd.concat([pd.Series([v] * len(arms[v][0])) for v in range(n)])
        result       = multivariate_logrank_test(all_t, all_g, all_e)
        p_global     = result.p_value
        p_global_str = "<0.001" if p_global < 0.001 else f"{p_global:.3f}"
        pair_lines   = []
        for (a, b) in combinations(range(n), 2):
            lr  = lr_test(arms[a][0], arms[b][0],
                          event_observed_A=arms[a][1], event_observed_B=arms[b][1])
            pp  = lr.p_value
            ps  = "<0.001" if pp < 0.001 else f"{pp:.3f}"
            pair_lines.append(f"{labels[a]} vs {labels[b]}: p={ps}")
        p_text = f"Global log-rank p = {p_global_str}\n" + "\n".join(pair_lines)

    # 200-day intervals: 0, 200, 400, … up to 1095
    risk_times = np.arange(0, THREE_YEARS + 1, risk_interval)

    # ── Figure layout ─────────────────────────────────────────────────────────
    # Risk table needs one row per group; give it enough vertical space so
    # rows don't overlap — 0.45 inches per group, minimum 1.2 inches total
    risk_panel_h = max(1.2, n * 0.55)
    fig_h = 6.5 + risk_panel_h

    fig = plt.figure(figsize=(11, fig_h))
    gs  = fig.add_gridspec(
        2, 1,
        height_ratios=[6.5, risk_panel_h],
        hspace=0.08,
    )
    ax_km   = fig.add_subplot(gs[0])
    ax_risk = fig.add_subplot(gs[1])

    # ── KM curves with 95% CI ─────────────────────────────────────────────────
    kmfs = {}
    for v in range(n):
        t, e = arms[v]
        kmf  = KaplanMeierFitter()
        kmf.fit(t, e, label=labels[v])
        kmf.plot_survival_function(
            ax=ax_km,
            ci_show=True,       # 95% CI shaded band
            color=GROUP_COLORS[v],
        )
        kmfs[v] = kmf

    ax_km.text(0.98, 0.98, p_text,
               transform=ax_km.transAxes, ha="right", va="top",
               fontsize=8.5,
               bbox=dict(boxstyle="round,pad=0.3", fc="white", ec="gray", alpha=0.8))
    ax_km.set_xlabel("")
    ax_km.set_ylabel("Survival Probability (95% CI)")
    ax_km.set_title("")
    ax_km.set_xlim(-30, THREE_YEARS)
    ax_km.set_xticks(risk_times)
    ax_km.set_xticklabels([str(int(t)) for t in risk_times])
    ax_km.legend(loc="lower left", fontsize=9)
    ax_km.set_ylim(0, 1.05)
    plt.setp(ax_km.get_xticklabels(), visible=False)

    # ── At-risk table ─────────────────────────────────────────────────────────
    # Evenly space group rows so they never overlap regardless of n
    y_positions = np.linspace(n - 1, 0, n)   # top group at highest y, bottom at 0

    ax_risk.set_yticks(y_positions)
    ax_risk.set_yticklabels(labels, fontsize=8.5, ha="right")
    ax_risk.yaxis.set_tick_params(length=0, pad=8)
    ax_risk.set_xlabel("Time (days)")
    ax_risk.set_xticks(risk_times)
    ax_risk.set_xticklabels([str(int(t)) for t in risk_times], fontsize=8.5)
    for spine in ["top", "right", "left"]:
        ax_risk.spines[spine].set_visible(False)
    ax_risk.spines["bottom"].set_visible(True)

    for y_pos, v in zip(y_positions, range(n)):
        for t in risk_times:
            n_at_risk = int((kmfs[v].durations >= t).sum())
            ax_risk.text(t, y_pos, str(n_at_risk),
                         ha="center", va="center",
                         fontsize=8.5, color=GROUP_COLORS[v])

    ax_risk.set_xlim(0, THREE_YEARS)
    # Add a little padding above/below so numbers don't touch the edges
    padding = 0.6
    ax_risk.set_ylim(-padding, n - 1 + padding)
    ax_risk.tick_params(axis="y", left=False)

    plt.tight_layout()
    fig.subplots_adjust(left=0.18)
    return fig, p_global_str


def plot_forest(df_cox, fig_label, outcome_label="Outcome"):
    df_c = df_cox.copy()
    hrs, los, his = [], [], []
    for _, row in df_c.iterrows():
        try:
            hrs.append(float(row["HR"]))
            lo, hi = str(row["95% CI"]).replace(" ", "").split("\u2013")
            los.append(float(lo)); his.append(float(hi))
        except Exception:
            hrs.append(np.nan); los.append(np.nan); his.append(np.nan)
    df_c["hr_val"] = hrs; df_c["lo_val"] = los; df_c["hi_val"] = his
    df_c = df_c.dropna(subset=["hr_val"])
    n_rows = len(df_c)
    fig, ax = plt.subplots(figsize=(10, max(4, n_rows * 0.55 + 2)))
    colors_f = ["#d62728" if is_significant(str(p)) else "#1f77b4" for p in df_c["P-Value"]]
    y_pos = list(range(n_rows, 0, -1))
    for i, (_, row) in enumerate(df_c.iterrows()):
        y = y_pos[i]
        ax.plot([row["lo_val"], row["hi_val"]], [y, y], color=colors_f[i], lw=1.5)
        ax.plot(row["hr_val"], y, "D", color=colors_f[i], markersize=7, zorder=3)
    ax.axvline(x=1, color="black", linestyle="--", lw=1, alpha=0.7)
    ax.set_yticks(y_pos)
    ax.set_yticklabels(df_c["Characteristic"].tolist(), fontsize=9)
    ax.set_xlabel("Hazard Ratio (95% CI)", fontsize=10); ax.set_title("")
    x_right = ax.get_xlim()[1]
    for i, (_, row) in enumerate(df_c.iterrows()):
        ax.text(x_right * 1.01, y_pos[i],
                f"{row['HR']} ({row['95% CI']})   p={row['P-Value']}",
                va="center", fontsize=8, color=colors_f[i])
    ax.legend(handles=[mpatches.Patch(color="#d62728", label="p < 0.05"),
                        mpatches.Patch(color="#1f77b4", label="p \u2265 0.05")],
              loc="lower right", fontsize=8)
    for spine in ["top", "right"]:
        ax.spines[spine].set_visible(False)
    plt.tight_layout()
    return fig




# ═════════════════════════════════════════════════════════════════════════════
# NARRATIVE HELPERS (rule-based — no external API required)
# ═════════════════════════════════════════════════════════════════════════════

def _date_range(df: pd.DataFrame) -> str:
    """Scan all columns for dates and return a year range string like '2015–2022'."""
    years = []
    for col in df.columns:
        parsed = pd.to_datetime(df[col], errors="coerce")
        valid = parsed.dropna()
        if len(valid) > len(df) * 0.1:
            years.extend(valid.dt.year.tolist())
    if not years:
        return ""
    lo, hi = int(min(years)), int(max(years))
    return f"{lo}–{hi}" if lo != hi else str(lo)


def _sig_rows(df_table: pd.DataFrame) -> pd.DataFrame:
    """Return only non-header rows with a significant p-value."""
    display_cols = [c for c in df_table.columns if c != "_section_header"]
    mask = df_table.get("_section_header", pd.Series(False, index=df_table.index)) == False
    data = df_table[mask][display_cols]
    return data[data["P-Value"].apply(lambda x: is_significant(str(x)))]


def _group_cols(df_table: pd.DataFrame, labels: list) -> list:
    """Return column names that correspond to group data (not Characteristic or P-Value)."""
    all_cols = [c for c in df_table.columns if c not in ("Characteristic", "P-Value", "_section_header")]
    return all_cols


def generate_study_population(df: pd.DataFrame, group_col: str, labels: list,
                               study_title: str, num_groups: int, date_range: str = "") -> str:
    ns = [(df[group_col] == v).sum() for v in range(num_groups)]
    total = sum(ns)
    group_parts = ", ".join(f"{lbl} (n={n:,})" for lbl, n in zip(labels, ns))
    date_sent = f" between {date_range}" if date_range else ""
    return (
        f"This study was conducted at Cleveland Clinic Main Campus{date_sent}. "
        f"A total of {total:,} patients were included in the analysis and stratified into "
        f"{num_groups} groups: {group_parts}. "
        f"Baseline characteristics, echocardiographic parameters, and clinical outcomes were "
        f"compared across groups using appropriate statistical methods, as detailed in the tables and figures below."
    )


def generate_table_narrative(table_name: str, df_table: pd.DataFrame, labels: list,
                              date_range: str = "", context: str = "") -> str:
    gcols = _group_cols(df_table, labels)
    sig = _sig_rows(df_table)
    display_cols = [c for c in df_table.columns if c != "_section_header"]
    mask = df_table.get("_section_header", pd.Series(False, index=df_table.index)) == False
    all_rows = df_table[mask][display_cols]
    n_total = len(all_rows)
    n_sig = len(sig)

    intro = f"{table_name} presents a comparison of {n_total} variables across the {len(labels)} study groups ({', '.join(labels)})."

    if n_sig == 0:
        body = "No statistically significant differences were observed between groups across the variables examined (all p\u22650.05)."
    else:
        findings = []
        for _, row in sig.iterrows():
            char = row.get("Characteristic", "")
            vals = " vs. ".join(str(row.get(c, "")) for c in gcols)
            p = row.get("P-Value", "")
            findings.append(f"{char} differed significantly across groups ({vals}; p={p})")
        if len(findings) == 1:
            body = f"A statistically significant difference was identified for {findings[0]}."
        elif len(findings) <= 4:
            body = "Statistically significant differences were observed for the following variables: " + "; ".join(findings) + "."
        else:
            # Summarise to avoid an overwhelming sentence
            top = "; ".join(findings[:4])
            body = (f"Significant between-group differences were identified for {n_sig} variables, including: "
                    f"{top}; among others (all p<0.05).")

    non_sig_n = n_total - n_sig
    close = f"The remaining {non_sig_n} variable(s) did not reach statistical significance (p\u22650.05)." if non_sig_n > 0 and n_sig > 0 else ""
    return " ".join(s for s in [intro, body, close] if s)


def generate_km_narrative(fig_label: str, km_time: str, km_event: str,
                           df: pd.DataFrame, group_col: str, labels: list,
                           km_p_str: str, date_range: str = "") -> str:
    THREE_YEARS = 1095
    ONE_YEAR = 365
    rates = {}
    for v, lbl in enumerate(labels):
        sub = df[df[group_col] == v].copy()
        t = pd.to_numeric(sub[km_time], errors="coerce")
        e = pd.to_numeric(sub[km_event], errors="coerce")
        idx = t.dropna().index.intersection(e.dropna().index)
        t, e = t.loc[idx], e.loc[idx]
        rates[lbl] = {}
        for tag, cutoff in [("1yr", ONE_YEAR), ("3yr", THREE_YEARS)]:
            e_c = e.where(t <= cutoff, other=0)
            t_c = t.clip(upper=cutoff)
            kmf = KaplanMeierFitter()
            try:
                kmf.fit(t_c, e_c)
                r = (1 - kmf.survival_function_at_times([cutoff]).values[0]) * 100
                rates[lbl][tag] = f"{r:.1f}%"
            except Exception:
                rates[lbl][tag] = "N/A"

    # Build readable rate strings
    rates_1yr = " vs. ".join(rates[lbl]["1yr"] for lbl in labels)
    rates_3yr = " vs. ".join(rates[lbl]["3yr"] for lbl in labels)

    sig_word = "statistically significant" if is_significant(km_p_str) else "not statistically significant"
    p_display = f"p={km_p_str}" if km_p_str != "<0.001" else "p<0.001"

    # Detect stepwise pattern (monotone increasing or decreasing 3-yr rates)
    vals_3yr = []
    for lbl in labels:
        try:
            vals_3yr.append(float(rates[lbl]["3yr"].replace("%", "")))
        except Exception:
            vals_3yr.append(None)
    valid_vals = [v for v in vals_3yr if v is not None]
    stepwise = ""
    if len(valid_vals) >= 3:
        if valid_vals == sorted(valid_vals):
            stepwise = " A stepwise increase in event rates was observed across the groups."
        elif valid_vals == sorted(valid_vals, reverse=True):
            stepwise = " A stepwise decrease in event rates was observed across the groups."

    short_title = fig_label.split(". ", 1)[-1] if ". " in fig_label else fig_label
    return (
        f"{fig_label} presents Kaplan-Meier event-free survival curves for {short_title} across the three study groups. "
        f"At 1 year, event rates were {rates_1yr} for {', '.join(labels)}, respectively. "
        f"At 3 years, event rates were {rates_3yr}.{stepwise} "
        f"The overall log-rank test was {sig_word} ({p_display})."
    )


def generate_cox_narrative(table_name: str, df_cox: pd.DataFrame, outcome_label: str,
                            labels: list, date_range: str = "") -> str:
    sig = df_cox[df_cox["P-Value"].apply(lambda x: is_significant(str(x)))]
    not_sig = df_cox[~df_cox["P-Value"].apply(lambda x: is_significant(str(x)))]

    intro = (f"Multivariate Cox proportional hazards regression was performed to identify independent "
             f"predictors of {outcome_label}.")

    if sig.empty:
        body = "No variables reached statistical significance in the multivariable model (all p\u22650.05)."
    else:
        parts = []
        for _, row in sig.iterrows():
            parts.append(f"{row['Characteristic']} (HR {row['HR']}, 95% CI {row['95% CI']}, p={row['P-Value']})")
        body = "The following variables were independently associated with " + outcome_label + ": " + "; ".join(parts) + "."

    if not not_sig.empty:
        ns_names = ", ".join(not_sig["Characteristic"].tolist())
        close = f"The following covariates did not reach statistical significance: {ns_names} (all p\u22650.05)."
    else:
        close = ""

    return " ".join(s for s in [intro, body, close] if s)


def generate_key_findings(
    study_title: str,
    labels: list,
    df_t1: pd.DataFrame,
    df_t2: pd.DataFrame,
    df_t3: pd.DataFrame,
    km_summaries: list,
    cox_summaries: list,
    date_range: str = "",
) -> str:
    findings = []
    idx = 1

    # 1. Population
    n_total = 0
    ns_str = ""
    # Try to infer from df_t1 column headers
    if not df_t1.empty:
        gcols = _group_cols(df_t1, labels)
        parts = []
        for c in gcols:
            # Header format is "Label (N=XXX)"
            m = re.search(r"N=(\d+)", c)
            if m:
                parts.append(f"{c.split(' (')[0]} (n={m.group(1)})")
                n_total += int(m.group(1))
        ns_str = ", ".join(parts)
    date_sent = f" between {date_range}" if date_range else ""
    if ns_str:
        findings.append(f"{idx}. This study{date_sent} included a total of {n_total:,} patients from Cleveland Clinic Main Campus, stratified into {len(labels)} groups: {ns_str}.")
    else:
        findings.append(f"{idx}. This study{date_sent} was conducted at Cleveland Clinic Main Campus and included patients stratified into {len(labels)} groups: {', '.join(labels)}.")
    idx += 1

    # 2. Significant baseline differences
    if not df_t1.empty:
        sig1 = _sig_rows(df_t1)
        if not sig1.empty:
            names = ", ".join(sig1["Characteristic"].tolist()[:5])
            findings.append(f"{idx}. Significant baseline differences were identified in the following variables: {names} (all p<0.05).")
            idx += 1
        else:
            findings.append(f"{idx}. Baseline characteristics were well-balanced across groups, with no statistically significant differences observed.")
            idx += 1

    # 3. Echo differences
    if not df_t2.empty:
        sig2 = _sig_rows(df_t2)
        if not sig2.empty:
            names = ", ".join(sig2["Characteristic"].tolist()[:4])
            gcols = _group_cols(df_t2, labels)
            row0 = sig2.iloc[0]
            vals = " vs. ".join(str(row0.get(c, "")) for c in gcols)
            findings.append(f"{idx}. Significant echocardiographic differences were observed, including {names}. For example, {row0['Characteristic']} was {vals} (p={row0['P-Value']}).")
            idx += 1

    # 4. Outcome differences
    if not df_t3.empty:
        sig3 = _sig_rows(df_t3)
        if not sig3.empty:
            gcols = _group_cols(df_t3, labels)
            for _, row in sig3.iterrows():
                vals = " vs. ".join(str(row.get(c, "")) for c in gcols)
                findings.append(f"{idx}. {row['Characteristic']} differed significantly across groups ({vals}; p={row['P-Value']}).")
                idx += 1
                if idx > 7:
                    break

    # 5. KM summaries
    for fig_label, p_str in km_summaries:
        sig_word = "significant" if is_significant(p_str) else "non-significant"
        short = fig_label.split(". ", 1)[-1] if ". " in fig_label else fig_label
        findings.append(f"{idx}. Kaplan-Meier analysis for {short} demonstrated a {sig_word} difference across groups (overall log-rank p={p_str}).")
        idx += 1

    # 6. Cox summaries
    for t_name, df_c, outcome in cox_summaries:
        if df_c is None or df_c.empty:
            continue
        sig_c = df_c[df_c["P-Value"].apply(lambda x: is_significant(str(x)))]
        if not sig_c.empty:
            parts = [f"{r['Characteristic']} (HR {r['HR']}, 95% CI {r['95% CI']}, p={r['P-Value']})" for _, r in sig_c.iterrows()]
            findings.append(f"{idx}. On multivariate Cox regression for {outcome}, independent predictors included: {'; '.join(parts)}.")
        else:
            findings.append(f"{idx}. Multivariate Cox regression for {outcome} identified no independently significant predictors.")
        idx += 1

    # Final clinical note
    findings.append(f"{idx}. These findings highlight meaningful differences across the {len(labels)} study groups and support further investigation into the clinical implications of the identified risk factors.")

    return "\n\n".join(findings)


def add_narrative_to_doc(doc, text: str):
    """Add a plain left-aligned narrative paragraph to the Word document."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = False
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    doc.add_paragraph()


# ═════════════════════════════════════════════════════════════════════════════
# STEP 4 — Variable selection
# ═════════════════════════════════════════════════════════════════════════════
st.header("4. Variable Selection")

resolved_cont_base    = resolve_vars(df, DEFAULT_CONT_BASE)
resolved_cat_base     = resolve_vars(df, DEFAULT_CAT_BASE)
resolved_echo         = resolve_vars(df, DEFAULT_ECHO)
resolved_outcomes_cat = resolve_vars(df, DEFAULT_OUTCOMES_CAT)
resolved_outcomes_cont= resolve_vars(df, DEFAULT_OUTCOMES_CONT)

def _store_key(section_label, lbl):
    """Persistent storage key — never used as a widget key, so Streamlit won't delete it."""
    safe = re.sub(r"[^a-z0-9]", "_", lbl.lower())[:50]
    return f"store_{section_label}__{safe}"

def _widget_key(section_label, lbl):
    """Widget key for the checkbox — separate from the storage key."""
    safe = re.sub(r"[^a-z0-9]", "_", lbl.lower())[:50]
    return f"chk_{section_label}__{safe}"

# Keep legacy alias so any external references still resolve
def _chk_key(section_label, lbl):
    return _store_key(section_label, lbl)

def _make_sync_callback(section_label, lbl):
    """Return an on_change callback that copies the widget value into the store key."""
    wk = _widget_key(section_label, lbl)
    sk = _store_key(section_label, lbl)
    def _cb():
        st.session_state[sk] = st.session_state[wk]
    return _cb

def var_checkboxes(section_label, resolved_dict, all_defaults):
    unmatched = [lbl for lbl in all_defaults if lbl not in resolved_dict]
    if unmatched:
        st.caption(f"Could not match in file (skipped): {', '.join(unmatched)}")
    if not resolved_dict:
        st.warning("No variables from this group were found in your file.")
        return {}

    # Initialise STORAGE keys to True only on first encounter.
    # These live under store_* keys which are never used as widget keys,
    # so Streamlit will never delete them when checkboxes are hidden by the search filter.
    for lbl in resolved_dict:
        sk = _store_key(section_label, lbl)
        if sk not in st.session_state:
            st.session_state[sk] = True

    # Search filter — only affects which checkboxes are VISIBLE, not their stored state
    search_term = st.text_input(
        "🔍 Search variables", key=f"search_{section_label}",
        placeholder="Type to filter..."
    )
    filtered = {lbl: col for lbl, col in resolved_dict.items()
                if search_term.lower() in lbl.lower() or search_term.lower() in col.lower()} \
               if search_term else resolved_dict

    if not filtered:
        st.info("No variables match your search.")
    else:
        c_sel, c_des, _ = st.columns([1, 1, 4])
        if c_sel.button("☑ Select all visible", key=f"selall_{section_label}"):
            for lbl in filtered:
                st.session_state[_store_key(section_label, lbl)] = True
        if c_des.button("☐ Deselect all visible", key=f"desall_{section_label}"):
            for lbl in filtered:
                st.session_state[_store_key(section_label, lbl)] = False

        grid = st.columns(3)
        for i, (lbl, col) in enumerate(filtered.items()):
            sk = _store_key(section_label, lbl)
            wk = _widget_key(section_label, lbl)
            # Seed the widget key from the store before rendering so the checkbox
            # shows the correct persisted value even after being hidden/unhidden.
            st.session_state[wk] = st.session_state[sk]
            grid[i % 3].checkbox(
                lbl,
                key=wk,
                on_change=_make_sync_callback(section_label, lbl),
                help=f"Matched to column: '{col}'",
            )

    # Collect from the FULL resolved_dict using the persistent store keys.
    return {lbl: col for lbl, col in resolved_dict.items()
            if st.session_state.get(_store_key(section_label, lbl), True)}

tab1, tab2, tab3 = st.tabs(["Table 1 \u2013 Baseline", "Table 2 \u2013 Echo", "Table 3 \u2013 Outcomes"])

with tab1:
    include_t1 = st.toggle("Include Table 1 in report", value=True)
    t1_name = st.text_input("Table 1 title", value="Table 1. Baseline Characteristics", key="t1name")
    if include_t1:
        st.subheader("Continuous variables")
        sel_cont_base = var_checkboxes("cont_base", resolved_cont_base, DEFAULT_CONT_BASE)
        st.subheader("Categorical variables")
        sel_cat_base  = var_checkboxes("cat_base", resolved_cat_base, DEFAULT_CAT_BASE)
    else:
        sel_cont_base = {}; sel_cat_base = {}

with tab2:
    include_t2 = st.toggle("Include Table 2 in report", value=True)
    t2_name = st.text_input("Table 2 title", value="Table 2. Echocardiographic Variables", key="t2name")
    sel_echo = var_checkboxes("echo", resolved_echo, DEFAULT_ECHO) if include_t2 else {}

with tab3:
    include_t3 = st.toggle("Include Table 3 in report", value=True)
    t3_name = st.text_input("Table 3 title", value="Table 3. Outcomes", key="t3name")
    if include_t3:
        st.subheader("Categorical outcomes")
        sel_outcomes_cat  = var_checkboxes("outcomes_cat",  resolved_outcomes_cat,  DEFAULT_OUTCOMES_CAT)
        st.subheader("Continuous outcomes")
        sel_outcomes_cont = var_checkboxes("outcomes_cont", resolved_outcomes_cont, DEFAULT_OUTCOMES_CONT)
    else:
        sel_outcomes_cat = {}; sel_outcomes_cont = {}

# ═════════════════════════════════════════════════════════════════════════════
# STEP 5 — KM curves (3-year, pre-defined outcomes)
# ═════════════════════════════════════════════════════════════════════════════
st.header("5. Kaplan-Meier Curves (3-Year Outcomes)")
st.caption(
    "Four standard outcomes are pre-defined for 3-year analysis. "
    "Toggle each on/off and map the time/event columns. "
    "If a column cannot be found it will be flagged below."
)

# Pre-defined outcome specs: (internal_key, default_display_title, time_keywords, event_keywords)
KM_OUTCOMES = [
    (
        "mortality",
        "3-Year All-Cause Mortality",
        ["Death_3year_days", "Death_1year_days", "Death_2year_days"],
        ["Death_3year_event", "Death_1year_event", "Death_2year_event"],
    ),
    (
        "hfhosp",
        "3-Year Heart Failure Hospitalization",
        ["HFH_3year_days", "HFH_Days", "HFH_1year_days", "HFH_2year_days"],
        ["HFH_3year_event", "HFH_1_0", "HFH_1year_event", "HFH_2year_event"],
    ),
    (
        "stroke",
        "3-Year Stroke",
        ["Stroke_3year_days", "Stroke_Days", "Stroke_1year_days", "Stroke_2year_days"],
        ["Stroke_3year_event", "Stroke_1_0", "Stroke_1year_event", "Stroke_2year_event"],
    ),
    (
        "ppm",
        "3-Year Permanent Pacemaker Implantation",
        ["PPM_3_DAYS", "PPM_Days"],
        ["PPM1_0_3_year", "PPM_1", "PPM_in_Less_than_30_Days"],
    ),
]

km_configs = []
for key, default_title, time_kws, event_kws in KM_OUTCOMES:
    auto_time  = find_column(df, time_kws)
    auto_event = find_column(df, event_kws)

    with st.expander(f"🫀 {default_title}", expanded=True):
        include_km = st.toggle(f"Include in report", value=True, key=f"km_include_{key}")
        if include_km:
            km_title = st.text_input("Figure title (editable)", value=default_title, key=f"km_title_{key}")
            c1, c2 = st.columns(2)

            if auto_time:
                time_idx = col_options.index(auto_time) + 1
                c1.success(f"Auto-matched time column: **{auto_time}**")
            else:
                time_idx = 0
                c1.warning(f"No {key.upper()} time column found — please select manually or leave as (none).")
            if auto_event:
                event_idx = col_options.index(auto_event) + 1
                c2.success(f"Auto-matched event column: **{auto_event}**")
            else:
                event_idx = 0
                c2.warning(f"No {key.upper()} event column found — please select manually or leave as (none).")

            km_time_options  = ["(none)"] + col_options
            km_event_options = ["(none)"] + col_options

            t_idx = km_time_options.index(auto_time)   if auto_time  and auto_time  in km_time_options  else 0
            e_idx = km_event_options.index(auto_event) if auto_event and auto_event in km_event_options else 0

            km_time  = c1.selectbox("Time column",        km_time_options,  index=t_idx, key=f"km_time_{key}")
            km_event = c2.selectbox("Event column (0/1)", km_event_options, index=e_idx, key=f"km_event_{key}")

            if km_time == "(none)" or km_event == "(none)":
                st.info(f"⚠️  No {'PPM' if key=='ppm' else key.upper()} data available — this KM curve will be skipped.")
            else:
                km_configs.append((km_title, km_time, km_event))

# ═════════════════════════════════════════════════════════════════════════════
# STEP 6 — Cox Proportional Hazards Analysis (all group counts)
# ═════════════════════════════════════════════════════════════════════════════
st.header("6. Cox Proportional Hazards Analysis (optional)")
st.caption(
    "Configure up to 4 Cox regression models. Select all covariates you want included. A forest plot is generated for each model."
)

run_cox = st.checkbox("Include Cox regression")

# When Cox is turned off, wipe all covariate selections so they start fresh next time
if not run_cox:
    stale = [k for k in st.session_state if k.startswith("cox_sel_")]
    for k in stale:
        del st.session_state[k]

cox_analyses = []

if run_cox:
    num_cox = st.number_input(
        "How many Cox regression analyses would you like to run?",
        min_value=1, max_value=4, value=1, step=1, key="num_cox"
    )
    st.caption(f"You have selected **{int(num_cox)}** Cox analysis model(s). Configure each one below.")
    st.divider()

    for ci in range(int(num_cox)):
        with st.expander(f"⚙️ Cox Analysis {ci+1}", expanded=True):
            t4_name_ci = st.text_input(
                "Table title",
                value="Cox Proportional Hazards Analysis",
                key=f"t4name_{ci}"
            )
            outcome_ci = st.text_input(
                "Outcome label (used in forest plot title)",
                value="Outcome", key=f"cox_outcome_{ci}"
            )
            c1, c2 = st.columns(2)
            cox_search_dur = c1.text_input("🔍 Filter time columns", value="",
                                            key=f"cox_sdur_{ci}", placeholder="Search...")
            cox_search_evt = c2.text_input("🔍 Filter event columns", value="",
                                            key=f"cox_sevt_{ci}", placeholder="Search...")
            filtered_dur = [c for c in col_options if cox_search_dur.lower() in c.lower()] \
                           if cox_search_dur else col_options
            filtered_evt = [c for c in col_options if cox_search_evt.lower() in c.lower()] \
                           if cox_search_evt else col_options
            dur_ci  = c1.selectbox("Duration / time column", filtered_dur, key=f"cox_dur_{ci}")
            evt_ci  = c2.selectbox("Event column (0/1)",     filtered_evt, key=f"cox_evt_{ci}")

            st.markdown("**Select covariates**")

            # Inline-filtered multiselect: type to narrow options, click to select.
            # Selections persist in session_state automatically via the widget key.
            cov_ci = st.multiselect(
                "Type to search & select covariates",
                options=col_options,
                default=[c for c in col_options
                         if st.session_state.get(f"cox_sel_{ci}_{c}", False)],
                key=f"cox_ms_{ci}",
                placeholder="Start typing a column name…"
            )
            # Sync selections back to session_state so they survive reruns
            for col in col_options:
                st.session_state[f"cox_sel_{ci}_{col}"] = (col in cov_ci)

            if cov_ci:
                st.caption(f"✅ {len(cov_ci)} covariate(s) selected: {', '.join(cov_ci)}")
            else:
                st.caption("No covariates selected yet.")

            cox_analyses.append({
                "t4_name":       t4_name_ci,
                "outcome_label": outcome_ci,
                "duration_col":  dur_ci,
                "event_col":     evt_ci,
                "covariates":    cov_ci,
            })

# Keep backward-compat single-analysis variables (used in run block)
cox_duration_col  = cox_analyses[0]["duration_col"]  if cox_analyses else None
cox_event_col     = cox_analyses[0]["event_col"]     if cox_analyses else None
cox_covariates    = cox_analyses[0]["covariates"]    if cox_analyses else []
cox_outcome_label = cox_analyses[0]["outcome_label"] if cox_analyses else "Outcome"
t4_name           = cox_analyses[0]["t4_name"]       if cox_analyses else "Cox Proportional Hazards Analysis"

def build_cox_df(df, group_col, duration_col, event_col, covariates):
    cols_needed = list(dict.fromkeys([duration_col, event_col] + covariates))
    missing = [c for c in cols_needed if c not in df.columns]
    if missing:
        return None, f"Missing columns: {', '.join(missing)}"
    cox_df = df[cols_needed].copy()
    # Coerce everything to numeric — whitespace strings and non-numeric values become NaN
    for c in cols_needed:
        cox_df[c] = pd.to_numeric(cox_df[c], errors="coerce")
    cox_df = cox_df.dropna()
    if len(cox_df) < 10:
        return None, "Not enough data for Cox regression (need ≥ 10 complete rows)."
    try:
        cph = CoxPHFitter()
        cph.fit(cox_df, duration_col=duration_col, event_col=event_col)
        rows = []
        for var in cph.summary.index:
            hr = np.exp(cph.summary.loc[var, "coef"])
            lo = np.exp(cph.summary.loc[var, "coef"] - 1.96 * cph.summary.loc[var, "se(coef)"])
            hi = np.exp(cph.summary.loc[var, "coef"] + 1.96 * cph.summary.loc[var, "se(coef)"])
            p  = cph.summary.loc[var, "p"]
            rows.append({"Characteristic": var, "HR": f"{hr:.2f}",
                         "95% CI": f"{lo:.2f} \u2013 {hi:.2f}", "P-Value": fmt_p(p)})
        return pd.DataFrame(rows), None
    except Exception as e:
        return None, str(e)

# ═════════════════════════════════════════════════════════════════════════════
# STEP 7 — Run
# ═════════════════════════════════════════════════════════════════════════════
st.divider()
if not st.button("\u25b6  Run analysis", type="primary"):
    st.stop()

fig_counter   = [1]
table_counter = [1]
def next_fig():
    n = fig_counter[0]; fig_counter[0] += 1; return n
def next_table():
    n = table_counter[0]; table_counter[0] += 1; return n

st.header(study_title)

# ── Compute date range from data ──────────────────────────────────────────────
date_range = _date_range(df)

# ── Study Population paragraph ────────────────────────────────────────────────
st.subheader("Study Population")
try:
    pop_narrative = generate_study_population(df, group_col, labels, study_title, num_groups, date_range)
except Exception as e:
    pop_narrative = f"[Study population narrative could not be generated: {e}]"
st.write(pop_narrative)
st.divider()

st.subheader("Study Flow")
flow_fig = plot_study_flow(df, group_col, labels, study_title)
st.pyplot(flow_fig)
f_num = next_fig()
st.caption(f"Figure {f_num}. Study flow diagram \u2013 {study_title}")
plt.close(flow_fig)
st.divider()

# ── Collect KM p-values for key findings (populated below) ───────────────────
km_p_registry = {}   # fig_label -> p_str

# Tables
df_t1 = pd.DataFrame()
if include_t1:
    t_num = next_table()
    t1_display = t1_name.replace("Table 1", f"Table {t_num}", 1) if t1_name.startswith("Table 1") else t1_name
    df_cont = build_stats_df(df, group_col, sel_cont_base, "continuous", labels)
    df_cat  = build_stats_df(df, group_col, sel_cat_base,  "categorical", labels)
    df_t1   = pd.concat([df_cont, df_cat], ignore_index=True)
    df_t1["_section_header"] = False
    df_t1 = inject_section_headers(df_t1, T1_SECTION_HEADERS)
    if df_t1.empty:
        st.warning("No variables found for Table 1.")
    else:
        try:
            t1_narrative = generate_table_narrative(t1_display, df_t1, labels, date_range)
        except Exception as e:
            t1_narrative = f"[Narrative could not be generated: {e}]"
        st.write(t1_narrative)
        st.subheader(t1_display)
        st.dataframe(style_pvalues(df_t1), use_container_width=True, hide_index=True)

df_t2 = pd.DataFrame()
if include_t2:
    t_num = next_table()
    t2_display = t2_name.replace("Table 2", f"Table {t_num}", 1) if t2_name.startswith("Table 2") else t2_name
    df_t2 = build_stats_df(df, group_col, sel_echo, "continuous", labels)
    df_t2["_section_header"] = False
    df_t2 = inject_section_headers(df_t2, T2_SECTION_HEADERS)
    if df_t2.empty:
        st.warning("No echo variables found.")
    else:
        try:
            t2_narrative = generate_table_narrative(t2_display, df_t2, labels, date_range)
        except Exception as e:
            t2_narrative = f"[Narrative could not be generated: {e}]"
        st.write(t2_narrative)
        st.subheader(t2_display)
        st.dataframe(style_pvalues(df_t2), use_container_width=True, hide_index=True)

df_t3 = pd.DataFrame()
if include_t3:
    t_num = next_table()
    t3_display = t3_name.replace("Table 3", f"Table {t_num}", 1) if t3_name.startswith("Table 3") else t3_name
    df_t3_cat  = build_stats_df(df, group_col, sel_outcomes_cat,  "categorical", labels)
    df_t3_cont = build_stats_df(df, group_col, sel_outcomes_cont, "continuous",  labels)
    df_t3 = pd.concat([df_t3_cat, df_t3_cont], ignore_index=True)
    df_t3["_section_header"] = False
    df_t3 = inject_section_headers(df_t3, T3_SECTION_HEADERS)
    if df_t3.empty:
        st.warning("No outcome variables found.")
    else:
        try:
            t3_narrative = generate_table_narrative(t3_display, df_t3, labels, date_range)
        except Exception as e:
            t3_narrative = f"[Narrative could not be generated: {e}]"
        st.write(t3_narrative)
        st.subheader(t3_display)
        st.dataframe(style_pvalues(df_t3), use_container_width=True, hide_index=True)

# KM curves
km_figs_for_export = []
if km_configs:
    st.subheader("Kaplan-Meier Curves")
    for km_short, km_time, km_event in km_configs:
        f_num = next_fig()
        fig_label = f"Figure {f_num}. {km_short}"
        try:
            fig, km_p_str = plot_km(df, group_col, km_time, km_event, labels, fig_label)
            km_p_registry[fig_label] = km_p_str
            try:
                km_narr = generate_km_narrative(fig_label, km_time, km_event, df, group_col, labels, km_p_str, date_range)
            except Exception as e:
                km_narr = f"[Narrative could not be generated: {e}]"
            st.write(km_narr)
            st.pyplot(fig)
            st.caption(fig_label)
            km_figs_for_export.append((fig_label, km_time, km_event, km_narr))
            plt.close(fig)
        except Exception as e:
            st.error(f"Could not generate KM curve: {e}")

# Cox — loop over all configured analyses
all_cox_results = []   # list of (t4_display, df_cox, forest_fig_label, forest_fig_for_export, outcome_label, cox_narr)
if run_cox and cox_analyses:
    for ca in cox_analyses:
        if not ca["duration_col"] or not ca["event_col"]:
            continue
        t_num = next_table()
        t4_display_ca = f"Table {t_num}. {ca['t4_name']}"
        extra_covs = list(dict.fromkeys(
            [c for c in ca["covariates"] if c not in (ca["duration_col"], ca["event_col"])]
        ))
        df_cox_ca, cox_error = build_cox_df(df, group_col, ca["duration_col"], ca["event_col"], extra_covs)
        if cox_error:
            st.error(cox_error)
            all_cox_results.append((t4_display_ca, None, "", None, ca["outcome_label"], ""))
        else:
            try:
                cox_narr = generate_cox_narrative(t4_display_ca, df_cox_ca, ca["outcome_label"], labels, date_range)
            except Exception as e:
                cox_narr = f"[Narrative could not be generated: {e}]"
            st.write(cox_narr)
            st.subheader(t4_display_ca)
            st.dataframe(style_pvalues(df_cox_ca), use_container_width=True, hide_index=True)
            f_num = next_fig()
            forest_label_ca = f"Figure {f_num}. Forest Plot \u2013 Multivariate Cox Regression for {ca['outcome_label']}"
            forest_fig_ca = plot_forest(df_cox_ca, forest_label_ca, ca["outcome_label"])
            st.pyplot(forest_fig_ca)
            st.caption(forest_label_ca)
            all_cox_results.append((t4_display_ca, df_cox_ca, forest_label_ca, forest_fig_ca, ca["outcome_label"], cox_narr))
            plt.close(forest_fig_ca)

# ── Key Findings Summary ──────────────────────────────────────────────────────
st.divider()
st.subheader("Key Findings Summary")
try:
    km_sums = [(fl, km_p_registry.get(fl, "N/A")) for fl, _, _, _ in km_figs_for_export]
    cox_sums = [(t, d, o) for t, d, _, _, o, _ in all_cox_results if d is not None]
    key_findings_text = generate_key_findings(
        study_title, labels, df_t1, df_t2, df_t3, km_sums, cox_sums, date_range
    )
except Exception as e:
    key_findings_text = f"[Key findings could not be generated: {e}]"
st.write(key_findings_text)

# Keep single-analysis backward-compat vars for the export block
df_cox = all_cox_results[0][1] if all_cox_results else None
t4_display = all_cox_results[0][0] if all_cox_results else t4_name
forest_fig_label = all_cox_results[0][2] if all_cox_results else ""
forest_fig_for_export = all_cox_results[0][3] if all_cox_results else None

# ═════════════════════════════════════════════════════════════════════════════
# STEP 8 — Export Word
# ═════════════════════════════════════════════════════════════════════════════
st.divider()
st.subheader("Export to Word")

doc = Document()
title_para = doc.add_heading(study_title, level=1)
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Study Population paragraph
doc.add_heading("Study Population", level=2)
try:
    add_narrative_to_doc(doc, pop_narrative)
except Exception:
    pass

try:
    flow_fig2 = plot_study_flow(df, group_col, labels, study_title)
    buf = io.BytesIO()
    flow_fig2.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    buf.seek(0)
    doc.add_picture(buf, width=Inches(6))
    cap = doc.add_paragraph(f"Figure 1. Study flow diagram \u2013 {study_title}")
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()
    plt.close(flow_fig2)
except Exception:
    pass

if not df_t1.empty:
    try:
        add_narrative_to_doc(doc, t1_narrative)
    except Exception:
        pass
    write_docx_table(doc, t1_name, df_t1)

if not df_t2.empty:
    try:
        add_narrative_to_doc(doc, t2_narrative)
    except Exception:
        pass
    write_docx_table(doc, t2_name, df_t2)

if not df_t3.empty:
    try:
        add_narrative_to_doc(doc, t3_narrative)
    except Exception:
        pass
    write_docx_table(doc, t3_name, df_t3)

for fig_label, km_time, km_event, km_narr in km_figs_for_export:
    try:
        add_narrative_to_doc(doc, km_narr)
    except Exception:
        pass
    try:
        fig2, _ = plot_km(df, group_col, km_time, km_event, labels, fig_label)
        buf = io.BytesIO()
        fig2.savefig(buf, format="png", dpi=150, bbox_inches="tight")
        buf.seek(0)
        doc.add_picture(buf, width=Inches(6))
        cap = doc.add_paragraph(fig_label)
        cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph()
        plt.close(fig2)
    except Exception:
        pass

for t4_disp_ex, df_cox_ex, forest_lbl_ex, _, outcome_ex, narr_ex in all_cox_results:
    if df_cox_ex is not None and not df_cox_ex.empty:
        try:
            add_narrative_to_doc(doc, narr_ex)
        except Exception:
            pass
        write_docx_table(doc, t4_disp_ex, df_cox_ex)
        try:
            forest_fig_ex2 = plot_forest(df_cox_ex, forest_lbl_ex, outcome_ex)
            buf = io.BytesIO()
            forest_fig_ex2.savefig(buf, format="png", dpi=150, bbox_inches="tight")
            buf.seek(0)
            doc.add_picture(buf, width=Inches(6))
            cap = doc.add_paragraph(forest_lbl_ex)
            cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
            doc.add_paragraph()
            plt.close(forest_fig_ex2)
        except Exception:
            pass

# Key Findings Summary at end of document
doc.add_page_break()
doc.add_heading("Key Findings Summary", level=1)
try:
    kf_para = doc.add_paragraph()
    kf_para.add_run(key_findings_text)
except Exception:
    pass

docx_buf = io.BytesIO()
doc.save(docx_buf)
docx_buf.seek(0)

st.download_button(
    label="\u2b07  Download Word report (.docx)",
    data=docx_buf,
    file_name=f"{safe_filename}.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
)